"""
Flask Backend — ExamGen AI with Google Drive Integration
"""

import os
import re
import json
import time
import glob
import uuid
import requests
import pdfkit
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import *

# Google Drive integration
try:
    from google_drive import upload_to_drive, list_drive_files, delete_drive_file, get_drive_file_info
    DRIVE_ENABLED = True
except Exception as e:
    print(f"Google Drive disabled: {e}")
    DRIVE_ENABLED = False

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAPER PATTERNS (same as before)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PATTERNS = {
    "chemistry": {
        "subject": "Chemistry",
        "total_marks": 60,
        "time_allowed": "2 Hours 30 Minutes",
        "sections": [
            {"section_name": "OBJECTIVE TYPE", "section_type": "MCQ", "question_label": "Q#1",
             "instructions": "Choose the correct answer. Each MCQ carries 1 mark.",
             "num_questions": 12, "marks_each": 1, "total_marks": 12, "attempt_rule": None},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#2",
             "instructions": "Attempt any FIVE (5) short questions out of 8. Each carries 2 marks.",
             "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#3",
             "instructions": "Attempt any FIVE (5) short questions out of 8. Each carries 2 marks.",
             "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#4",
             "instructions": "Attempt any FIVE (5) short questions out of 8. Each carries 2 marks.",
             "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#5",
             "instructions": "Note: Attempt any TWO (2) questions from Q#5, Q#6, Q#7.",
             "num_questions": 1, "total_marks": 9,
             "sub_parts": [{"part": "a", "marks": 5, "type": "descriptive"}, {"part": "b", "marks": 4, "type": "descriptive"}],
             "attempt_rule": "Attempt any 2 out of 3 (Q#5, Q#6, Q#7)"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#6",
             "instructions": "", "num_questions": 1, "total_marks": 9,
             "sub_parts": [{"part": "a", "marks": 5, "type": "descriptive"}, {"part": "b", "marks": 4, "type": "descriptive"}],
             "attempt_rule": "Attempt any 2 out of 3 (Q#5, Q#6, Q#7)"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#7",
             "instructions": "", "num_questions": 1, "total_marks": 9,
             "sub_parts": [{"part": "a", "marks": 5, "type": "descriptive"}, {"part": "b", "marks": 4, "type": "descriptive"}],
             "attempt_rule": "Attempt any 2 out of 3 (Q#5, Q#6, Q#7)"},
        ],
    },
    "biology": {
        "subject": "Biology", "total_marks": 60, "time_allowed": "2 Hours 30 Minutes",
        "sections": [
            {"section_name": "OBJECTIVE TYPE", "section_type": "MCQ", "question_label": "Q#1",
             "instructions": "Choose the correct answer.", "num_questions": 12, "marks_each": 1, "total_marks": 12, "attempt_rule": None},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#2",
             "instructions": "Attempt any 5 out of 8.", "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#3",
             "instructions": "Attempt any 5 out of 8.", "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#4",
             "instructions": "Attempt any 5 out of 8.", "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#5",
             "instructions": "Attempt any 2 out of 3.", "num_questions": 1, "total_marks": 9,
             "sub_parts": [{"part": "a", "marks": 5}, {"part": "b", "marks": 4}], "attempt_rule": "Attempt any 2 out of 3"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#6",
             "num_questions": 1, "total_marks": 9, "sub_parts": [{"part": "a", "marks": 5}, {"part": "b", "marks": 4}], "attempt_rule": "Attempt any 2 out of 3"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#7",
             "num_questions": 1, "total_marks": 9, "sub_parts": [{"part": "a", "marks": 5}, {"part": "b", "marks": 4}], "attempt_rule": "Attempt any 2 out of 3"},
        ],
    },
    "physics": {
        "subject": "Physics", "total_marks": 60, "time_allowed": "2 Hours 30 Minutes",
        "sections": [
            {"section_name": "OBJECTIVE TYPE", "section_type": "MCQ", "question_label": "Q#1",
             "instructions": "Choose the correct answer.", "num_questions": 12, "marks_each": 1, "total_marks": 12, "attempt_rule": None},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#2",
             "instructions": "Attempt any 5 out of 8.", "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#3",
             "instructions": "Attempt any 5 out of 8.", "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-I)", "section_type": "SHORT", "question_label": "Q#4",
             "instructions": "Attempt any 5 out of 8.", "num_questions": 8, "marks_each": 2, "total_marks": 10, "attempt_rule": "Attempt any 5 out of 8"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#5",
             "instructions": "Attempt any 2 out of 3.", "num_questions": 1, "total_marks": 9,
             "sub_parts": [{"part": "a", "marks": 4, "type": "descriptive"}, {"part": "b", "marks": 5, "type": "numerical"}], "attempt_rule": "Attempt any 2 out of 3"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#6",
             "num_questions": 1, "total_marks": 9, "sub_parts": [{"part": "a", "marks": 4}, {"part": "b", "marks": 5, "type": "numerical"}], "attempt_rule": "Attempt any 2 out of 3"},
            {"section_name": "SUBJECTIVE TYPE (Part-II)", "section_type": "LONG", "question_label": "Q#7",
             "num_questions": 1, "total_marks": 9, "sub_parts": [{"part": "a", "marks": 4}, {"part": "b", "marks": 5, "type": "numerical"}], "attempt_rule": "Attempt any 2 out of 3"},
        ],
    },
    "english": {
        "subject": "English", "total_marks": 75, "time_allowed": "2 Hours 30 Minutes",
        "sections": [
            {"section_name": "OBJECTIVE PAPER", "section_type": "MCQ_MIXED", "question_label": "Q#1",
             "instructions": "Choose the correct option.", "total_marks": 19, "attempt_rule": None,
             "sub_sections": [{"name": "Correct Form of Verb", "num_questions": 5}, {"name": "Spellings", "num_questions": 4},
                              {"name": "Meanings", "num_questions": 5}, {"name": "Grammar", "num_questions": 5}]},
            {"section_name": "SUBJECTIVE PAPER", "section_type": "MIXED", "question_label": "Q#2-Q#9",
             "instructions": "Answer as directed.", "total_marks": 56, "attempt_rule": None},
        ],
    },
    "maths": {
        "subject": "Mathematics", "total_marks": 75, "time_allowed": "2 Hours 30 Minutes",
        "sections": [
            {"section_name": "OBJECTIVE (MCQs)", "section_type": "MCQ", "question_label": "Q#1",
             "instructions": "Choose the correct answer.", "num_questions": 15, "marks_each": 1, "total_marks": 15, "attempt_rule": None},
            {"section_name": "SUBJECTIVE (Short)", "section_type": "SHORT", "question_label": "Q#2",
             "instructions": "Solve any 6 out of 9.", "num_questions": 9, "marks_each": 2, "total_marks": 12, "attempt_rule": "Solve any 6 out of 9"},
            {"section_name": "SUBJECTIVE (Short)", "section_type": "SHORT", "question_label": "Q#3",
             "instructions": "Solve any 6 out of 9.", "num_questions": 9, "marks_each": 2, "total_marks": 12, "attempt_rule": "Solve any 6 out of 9"},
            {"section_name": "SUBJECTIVE (Short)", "section_type": "SHORT", "question_label": "Q#4",
             "instructions": "Solve any 6 out of 9.", "num_questions": 9, "marks_each": 2, "total_marks": 12, "attempt_rule": "Solve any 6 out of 9"},
            {"section_name": "SUBJECTIVE (Long)", "section_type": "LONG", "question_label": "Q#5-Q#9",
             "instructions": "Attempt any 3 out of 5.", "num_questions": 5, "total_marks": 24,
             "sub_parts": [{"part": "a", "marks": 4}, {"part": "b", "marks": 4}], "attempt_rule": "Attempt any 3 out of 5"},
        ],
    },
}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# HELPERS (same as before, abbreviated)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_session_id():
    return str(uuid.uuid4())[:8]

def cleanup_old_files():
    now = time.time()
    max_age = AUTO_DELETE_DAYS * 86400
    for folder in [UPLOAD_FOLDER, os.path.join(OUTPUT_FOLDER, "pdf"),
                   os.path.join(OUTPUT_FOLDER, "docx"), os.path.join(OUTPUT_FOLDER, "json")]:
        if not os.path.exists(folder):
            continue
        for f in glob.glob(os.path.join(folder, "*")):
            if os.path.isfile(f) and (now - os.path.getmtime(f)) > max_age:
                os.remove(f)

def call_llm(system_msg, user_msg, max_tokens=8192):
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {A4F_API_KEY}"}
    payload = {"model": A4F_MODEL, "messages": [{"role": "system", "content": system_msg}, {"role": "user", "content": user_msg}],
               "max_tokens": max_tokens, "temperature": 0.3}
    resp = requests.post(A4F_API_URL, headers=headers, json=payload, timeout=180)
    if resp.status_code != 200:
        raise RuntimeError(f"API error {resp.status_code}")
    data = resp.json()
    raw = data.get("choices", [{}])[0].get("message", {}).get("content", "") or data.get("content", "")
    if "<think>" in raw:
        idx = raw.find("</think>")
        if idx != -1:
            raw = raw[idx + 8:]
    return raw.strip()

def extract_json(text):
    if not text:
        return None
    try:
        return json.loads(text.strip())
    except:
        pass
    for pat in (r"```json\s*([\s\S]*?)\s*```", r"```\s*([\s\S]*?)\s*```"):
        for m in re.findall(pat, text):
            try:
                return json.loads(m.strip())
            except:
                continue
    start = text.find("{")
    if start != -1:
        depth = 0
        for i in range(start, len(text)):
            if text[i] == "{": depth += 1
            elif text[i] == "}":
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[start:i + 1])
                    except:
                        break
    return None

def ocr_extract(filepath):
    payload = {"apikey": OCR_SPACE_API_KEY, "language": "eng", "OCREngine": "1", "isTable": "true", "scale": "true"}
    with open(filepath, "rb") as f:
        resp = requests.post(OCR_API_URL, files={"file": (os.path.basename(filepath), f)}, data=payload, timeout=60)
    if resp.status_code != 200:
        return ""
    result = resp.json()
    if result.get("OCRExitCode", 0) != 1:
        return ""
    return "\n".join(p.get("ParsedText", "") for p in result.get("ParsedResults", [])).strip()

def local_clean(text):
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        line = line.strip()
        if not line:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        line = "".join(c for c in line if not (0x0600 <= ord(c) <= 0x06FF or 0xFB50 <= ord(c) <= 0xFDFF))
        if any(x in line.upper() for x in ["WWW.", "FREEILM", ".COM"]):
            continue
        if line.count("@") > 1:
            continue
        line = re.sub(r'^(\d+)\,(\s)', r'\1.\2', line)
        line = re.sub(r'(\d),(\d)', r'\1.\2', line)
        line = re.sub(r'\bofthe\b', 'of the', line)
        line = re.sub(r'(\d+)grams?\b', r'\1 grams', line)
        line = re.sub(r'\[\s*(\d+)\s*\]', r'[\1 marks]', line)
        line = re.sub(r'\t+', '  ', line).strip()
        if len(line) >= 2:
            cleaned.append(line)
    return "\n".join(cleaned).strip()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PDF & DOCX GENERATORS (same as before)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def build_exam_html(exam):
    title = exam.get("exam_title", "Examination")
    subject = exam.get("subject", "")
    marks = exam.get("total_marks", "")
    time_a = exam.get("time_allowed", "")

    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>
@page {{ size: A4; margin: 10mm 12mm; }}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: 'Times New Roman', serif; font-size: 10pt; line-height: 1.35; color: #000; }}
.page {{ width: 100%; page-break-after: always; }}
.page:last-child {{ page-break-after: avoid; }}
.header {{ text-align: center; border-bottom: 2.5px double #000; padding-bottom: 6px; margin-bottom: 6px; }}
.academy {{ font-size: 16pt; font-weight: bold; letter-spacing: 3px; text-transform: uppercase; }}
.exam-title {{ font-size: 12pt; margin: 2px 0; }}
.subject-line {{ font-size: 11pt; font-weight: bold; }}
.meta {{ display: flex; justify-content: space-between; font-size: 9.5pt; margin: 4px 0; font-weight: bold; }}
.student-row {{ display: flex; justify-content: space-between; font-size: 9.5pt; margin: 3px 0 8px; }}
.student-row span {{ border-bottom: 1px solid #666; min-width: 150px; display: inline-block; }}
.sec-h {{ font-size: 10.5pt; font-weight: bold; text-transform: uppercase; background: #f0f0f0; padding: 3px 8px; margin: 8px 0 4px; border-left: 3px solid #000; }}
.q-label {{ font-size: 10pt; font-weight: bold; text-decoration: underline; margin: 6px 0 2px; }}
.inst {{ font-style: italic; font-size: 9pt; color: #333; margin: 1px 0 4px 12px; }}
.rule {{ font-weight: bold; color: #900; font-size: 9pt; margin: 1px 0 4px 12px; }}
.q {{ margin: 3px 0; font-size: 9.5pt; }}
.q-num {{ font-weight: bold; }}
.q-marks {{ float: right; font-size: 8.5pt; color: #555; font-weight: bold; }}
.opts {{ margin: 2px 0 4px 18px; display: grid; grid-template-columns: 1fr 1fr; gap: 1px 15px; font-size: 9.5pt; }}
.sub {{ margin: 2px 0 2px 18px; font-size: 9.5pt; }}
.sub-l {{ font-weight: bold; }}
.footer {{ text-align: center; margin-top: 8px; padding-top: 4px; border-top: 2px solid #000; font-size: 9pt; font-weight: bold; letter-spacing: 2px; }}
</style></head><body>"""

    sections_html = ""
    current_section = ""
    for sec in exam.get("sections", []):
        sn = sec.get("section_name", "")
        if sn != current_section:
            sections_html += f'<div class="sec-h">{sn}</div>\n'
            current_section = sn
        sections_html += f'<div class="q-label">{sec.get("question_label", "")}:</div>\n'
        if sec.get("instructions"):
            sections_html += f'<div class="inst">{sec["instructions"]}</div>\n'
        if sec.get("attempt_rule"):
            sections_html += f'<div class="rule">Note: {sec["attempt_rule"]}</div>\n'
        for q in sec.get("questions", []):
            st = sec.get("section_type", "")
            if st in ("MCQ", "MCQ_MIXED"):
                sections_html += f'<div class="q"><span class="q-num">({q.get("question_number","")})</span> {q.get("question_text","")}\n'
                if q.get("options"):
                    sections_html += '<div class="opts">'
                    for L in "ABCD":
                        if q["options"].get(L):
                            sections_html += f'<div><span class="q-num">({L})</span> {q["options"][L]}</div>'
                    sections_html += '</div>'
                sections_html += '</div>\n'
            elif q.get("sub_parts"):
                sections_html += f'<div class="q"><span class="q-num">{sec.get("question_label","")}</span> <span class="q-marks">[{q.get("marks","")} Marks]</span></div>\n'
                for sp in q["sub_parts"]:
                    sections_html += f'<div class="sub"><span class="sub-l">({sp.get("part","")})</span> {sp.get("text","")} [{sp.get("marks","")}]</div>\n'
            else:
                sections_html += f'<div class="q"><span class="q-num">({q.get("question_number","")})</span> {q.get("question_text","")} <span class="q-marks">[{q.get("marks","")}]</span></div>\n'

    all_lines = sections_html.strip().split("\n")
    mid = len(all_lines) // 2
    for i in range(mid, min(mid + 10, len(all_lines))):
        if 'sec-h' in all_lines[i]:
            mid = i
            break

    page_header = f"""<div class="header"><div class="academy">{ACADEMY_NAME}</div><div class="exam-title">{title}</div><div class="subject-line">Subject: {subject}</div></div>
<div class="meta"><div>Total Marks: {marks}</div><div>Time: {time_a}</div></div>
<div class="student-row"><div>Name: <span></span></div><div>Roll No: <span></span></div></div>"""

    html += f'<div class="page">{page_header}\n{"".join(all_lines[:mid])}</div>\n'
    html += f'<div class="page">{page_header}\n{"".join(all_lines[mid:])}\n<div class="footer">✦ END OF PAPER ✦</div></div>\n'
    html += "</body></html>"
    return html

def generate_pdf(exam, session_id):
    html = build_exam_html(exam)
    pdf_dir = os.path.join(OUTPUT_FOLDER, "pdf")
    os.makedirs(pdf_dir, exist_ok=True)
    html_path = os.path.join(pdf_dir, f"{session_id}.html")
    pdf_path = os.path.join(pdf_dir, f"{session_id}.pdf")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    opts = {"page-size": "A4", "margin-top": "0mm", "margin-bottom": "0mm", "margin-left": "0mm", "margin-right": "0mm", "encoding": "UTF-8", "no-outline": None}
    pdfkit.from_file(html_path, pdf_path, options=opts)
    return pdf_path

def generate_docx(exam, session_id):
    docx_dir = os.path.join(OUTPUT_FOLDER, "docx")
    os.makedirs(docx_dir, exist_ok=True)
    docx_path = os.path.join(docx_dir, f"{session_id}.docx")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = Cm(1)
        s.bottom_margin = Cm(1)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ACADEMY_NAME)
    r.bold = True
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(exam.get("exam_title", "")).font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Subject: {exam.get('subject', '')}")
    r.bold = True
    r.font.size = Pt(12)
    current_section = ""
    for sec in exam.get("sections", []):
        sn = sec.get("section_name", "")
        if sn != current_section:
            p = doc.add_paragraph()
            r = p.add_run(sn.upper())
            r.bold = True
            r.underline = True
            current_section = sn
        p = doc.add_paragraph()
        r = p.add_run(f"{sec.get('question_label', '')}:")
        r.bold = True
        r.underline = True
        for q in sec.get("questions", []):
            p = doc.add_paragraph()
            r = p.add_run(f"({q.get('question_number', '')}) ")
            r.bold = True
            p.add_run(q.get("question_text", ""))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("✦ END OF PAPER ✦").bold = True
    doc.save(docx_path)
    return docx_path


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.route("/")
def index():
    cleanup_old_files()
    return render_template("index.html")


@app.route("/api/subjects", methods=["GET"])
def get_subjects():
    subjects = [{"id": k, "name": v["subject"], "total_marks": v["total_marks"]} for k, v in PATTERNS.items()]
    return jsonify({"subjects": subjects})


@app.route("/api/upload", methods=["POST"])
def upload_images():
    if "images" not in request.files:
        return jsonify({"error": "No images"}), 400
    files = request.files.getlist("images")
    session_id = generate_session_id()
    session_dir = os.path.join(UPLOAD_FOLDER, session_id)
    os.makedirs(session_dir, exist_ok=True)
    all_text = ""
    for f in files:
        if f and allowed_file(f.filename):
            filepath = os.path.join(session_dir, secure_filename(f.filename))
            f.save(filepath)
            text = ocr_extract(filepath)
            if text:
                all_text += text + "\n\n"
    if not all_text.strip():
        return jsonify({"error": "Could not extract text"}), 400
    return jsonify({"session_id": session_id, "raw_text": all_text.strip(), "word_count": len(all_text.split())})


@app.route("/api/clean", methods=["POST"])
def clean_text():
    data = request.get_json()
    raw_text = data.get("raw_text", "")
    subject = data.get("subject", "General")
    try:
        system = "You are an OCR text fixer. Fix spelling, remove garbage, keep clean English. Output only cleaned text."
        cleaned = call_llm(system, f"Subject: {subject}\n\nFix:\n\n{raw_text}", 4096)
    except:
        cleaned = local_clean(raw_text)
    return jsonify({"cleaned_text": cleaned, "word_count": len(cleaned.split())})


@app.route("/api/generate", methods=["POST"])
def generate_exam():
    data = request.get_json()
    cleaned_text = data.get("cleaned_text", "")
    subject_id = data.get("subject", "chemistry")
    session_id = data.get("session_id", generate_session_id())
    
    pattern = PATTERNS.get(subject_id)
    if not pattern:
        return jsonify({"error": "Invalid subject"}), 400
    
    if not cleaned_text:
        return jsonify({"error": "No text provided"}), 400

    # Build section description for prompt
    section_desc = ""
    for sec in pattern["sections"]:
        label = sec.get("question_label", "")
        stype = sec.get("section_type", "")
        num_q = sec.get("num_questions", 0)
        marks_each = sec.get("marks_each", 0)
        total_m = sec.get("total_marks", 0)
        rule = sec.get("attempt_rule", "")
        
        section_desc += f"\n{label}: {sec.get('section_name', '')}"
        section_desc += f"\n  Type: {stype}, Questions: {num_q}"
        if marks_each:
            section_desc += f", {marks_each} marks each"
        section_desc += f", Total: {total_m} marks"
        if rule:
            section_desc += f"\n  Rule: {rule}"
        if "sub_parts" in sec:
            for sp in sec["sub_parts"]:
                section_desc += f"\n  Part ({sp['part']}): {sp['marks']} marks"

    system = """You are an expert exam paper generator. Output ONLY valid JSON, no explanation.

CRITICAL: Every question MUST have these fields:
- question_number (integer)
- question_text (string, never empty)
- marks (integer)

MCQ questions MUST also have:
- options: {"A": "...", "B": "...", "C": "...", "D": "..."}
- correct_answer: "A", "B", "C", or "D"

LONG questions with sub_parts MUST have:
- sub_parts: [{"part": "a", "text": "...", "marks": 5}, {"part": "b", "text": "...", "marks": 4}]"""

    user = f"""Generate a {pattern['subject']} exam paper.

Total Marks: {pattern['total_marks']}
Time: {pattern['time_allowed']}

PAPER PATTERN:
{section_desc}

STUDY MATERIAL:
{cleaned_text}

OUTPUT THIS EXACT JSON STRUCTURE:
{{
  "exam_title": "Annual Examination",
  "subject": "{pattern['subject']}",
  "total_marks": {pattern['total_marks']},
  "time_allowed": "{pattern['time_allowed']}",
  "sections": [
    {{
      "question_label": "Q#1",
      "section_name": "OBJECTIVE TYPE",
      "section_type": "MCQ",
      "instructions": "Choose the correct answer. Each MCQ carries 1 mark.",
      "attempt_rule": null,
      "questions": [
        {{
          "question_number": 1,
          "question_text": "What is the atomic number of Carbon?",
          "options": {{"A": "6", "B": "8", "C": "12", "D": "14"}},
          "correct_answer": "A",
          "marks": 1
        }}
      ]
    }},
    {{
      "question_label": "Q#2",
      "section_name": "SUBJECTIVE TYPE (Part-I)",
      "section_type": "SHORT",
      "instructions": "Attempt any 5 out of 8. Each carries 2 marks.",
      "attempt_rule": "Attempt any 5 out of 8",
      "questions": [
        {{
          "question_number": 1,
          "question_text": "Define atomic mass.",
          "marks": 2
        }}
      ]
    }},
    {{
      "question_label": "Q#5",
      "section_name": "SUBJECTIVE TYPE (Part-II)",
      "section_type": "LONG",
      "instructions": "Attempt any 2 out of 3.",
      "attempt_rule": "Attempt any 2 out of 3",
      "questions": [
        {{
          "question_number": 1,
          "question_text": "Explain the following:",
          "sub_parts": [
            {{"part": "a", "text": "Explain ionic bonding with examples.", "marks": 5}},
            {{"part": "b", "text": "Differentiate between ionic and covalent bonds.", "marks": 4}}
          ],
          "marks": 9
        }}
      ]
    }}
  ]
}}

Generate the complete exam now with ALL sections filled:"""

    try:
        raw = call_llm(system, user, max_tokens=8192)
        exam = extract_json(raw)
        
        if not exam:
            return jsonify({"error": "Failed to parse exam JSON from AI response"}), 500
        
        # Validate and fix the exam structure
        exam = validate_and_fix_exam(exam, pattern)
        
    except Exception as e:
        return jsonify({"error": f"AI generation failed: {str(e)}"}), 500

    # Save JSON
    json_dir = os.path.join(OUTPUT_FOLDER, "json")
    os.makedirs(json_dir, exist_ok=True)
    with open(os.path.join(json_dir, f"{session_id}.json"), "w") as f:
        json.dump(exam, f, indent=2)

    return jsonify({"session_id": session_id, "exam": exam})


def validate_and_fix_exam(exam, pattern):
    """Validate and fix the exam JSON structure to prevent undefined values."""
    
    # Ensure top-level fields exist
    if not exam.get("exam_title"):
        exam["exam_title"] = "Annual Examination"
    if not exam.get("subject"):
        exam["subject"] = pattern["subject"]
    if not exam.get("total_marks"):
        exam["total_marks"] = pattern["total_marks"]
    if not exam.get("time_allowed"):
        exam["time_allowed"] = pattern["time_allowed"]
    
    # Ensure sections exist
    if not exam.get("sections") or not isinstance(exam["sections"], list):
        exam["sections"] = []
    
    # Fix each section
    for i, sec in enumerate(exam["sections"]):
        # Ensure section fields
        if not sec.get("section_name"):
            sec["section_name"] = f"Section {i+1}"
        if not sec.get("question_label"):
            sec["question_label"] = f"Q#{i+1}"
        if not sec.get("section_type"):
            sec["section_type"] = "SHORT"
        if not sec.get("instructions"):
            sec["instructions"] = ""
        if sec.get("attempt_rule") is None:
            sec["attempt_rule"] = None
        
        # Ensure questions exist
        if not sec.get("questions") or not isinstance(sec["questions"], list):
            sec["questions"] = []
        
        # Fix each question
        for j, q in enumerate(sec["questions"]):
            # Ensure question fields
            if not q.get("question_number"):
                q["question_number"] = j + 1
            if not q.get("question_text"):
                q["question_text"] = f"Question {j+1}"
            if not q.get("marks"):
                q["marks"] = 1
            
            # Fix MCQ specific fields
            if sec["section_type"] == "MCQ":
                if not q.get("options") or not isinstance(q["options"], dict):
                    q["options"] = {"A": "Option A", "B": "Option B", "C": "Option C", "D": "Option D"}
                else:
                    # Ensure all options exist
                    for letter in ["A", "B", "C", "D"]:
                        if not q["options"].get(letter):
                            q["options"][letter] = f"Option {letter}"
                
                if not q.get("correct_answer") or q["correct_answer"] not in ["A", "B", "C", "D"]:
                    q["correct_answer"] = "A"
            
            # Fix LONG question sub_parts
            if sec["section_type"] == "LONG" and q.get("sub_parts"):
                if not isinstance(q["sub_parts"], list):
                    q["sub_parts"] = []
                
                for sp in q["sub_parts"]:
                    if not sp.get("part"):
                        sp["part"] = "a"
                    if not sp.get("text"):
                        sp["text"] = "Sub-question"
                    if not sp.get("marks"):
                        sp["marks"] = 4
    
    return exam


@app.route("/api/download/pdf", methods=["POST"])
def download_pdf():
    data = request.get_json()
    exam = data.get("exam")
    session_id = data.get("session_id", generate_session_id())
    if not exam:
        return jsonify({"error": "No exam data"}), 400
    try:
        pdf_path = generate_pdf(exam, session_id)
        return send_file(pdf_path, as_attachment=True, download_name=f"ghori_academy_{exam.get('subject','exam').lower()}_exam.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/download/docx", methods=["POST"])
def download_docx():
    data = request.get_json()
    exam = data.get("exam")
    session_id = data.get("session_id", generate_session_id())
    if not exam:
        return jsonify({"error": "No exam data"}), 400
    try:
        docx_path = generate_docx(exam, session_id)
        return send_file(docx_path, as_attachment=True, download_name=f"ghori_academy_{exam.get('subject','exam').lower()}_exam.docx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# GOOGLE DRIVE ROUTES (NEW)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.route("/api/drive/status", methods=["GET"])
def drive_status():
    """Check if Google Drive is enabled and configured."""
    return jsonify({
        "enabled": DRIVE_ENABLED,
        "folder_configured": bool(GOOGLE_DRIVE_FOLDER_ID),
    })


@app.route("/api/drive/upload", methods=["POST"])
def drive_upload():
    """Upload PDF to Google Drive with custom name."""
    if not DRIVE_ENABLED:
        return jsonify({"error": "Google Drive not configured"}), 400
    
    data = request.get_json()
    exam = data.get("exam")
    session_id = data.get("session_id", generate_session_id())
    custom_name = data.get("custom_name", "").strip()
    file_type = data.get("file_type", "pdf")  # 'pdf' or 'docx'
    
    if not exam:
        return jsonify({"error": "No exam data"}), 400
    
    try:
        # Generate file locally first
        if file_type == "docx":
            file_path = generate_docx(exam, session_id)
        else:
            file_path = generate_pdf(exam, session_id)
        
        # Generate default name if not provided
        if not custom_name:
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            custom_name = f"{ACADEMY_NAME}_{exam.get('subject', 'Exam')}_{timestamp}"
        
        # Upload to Drive
        result = upload_to_drive(file_path, custom_name, file_type)
        
        return jsonify({
            "success": True,
            "file": result,
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/drive/files", methods=["GET"])
def drive_files():
    """List all files in Google Drive folder."""
    if not DRIVE_ENABLED:
        return jsonify({"error": "Google Drive not configured", "files": []}), 200
    
    try:
        files = list_drive_files()
        return jsonify({"files": files})
    except Exception as e:
        return jsonify({"error": str(e), "files": []}), 200


@app.route("/api/drive/delete/<file_id>", methods=["DELETE"])
def drive_delete(file_id):
    """Delete a file from Google Drive."""
    if not DRIVE_ENABLED:
        return jsonify({"error": "Google Drive not configured"}), 400
    
    try:
        success = delete_drive_file(file_id)
        if success:
            return jsonify({"success": True})
        else:
            return jsonify({"error": "Could not delete file"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(os.path.join(OUTPUT_FOLDER, "pdf"), exist_ok=True)
    os.makedirs(os.path.join(OUTPUT_FOLDER, "docx"), exist_ok=True)
    os.makedirs(os.path.join(OUTPUT_FOLDER, "json"), exist_ok=True)
    os.makedirs(CREDENTIALS_FOLDER, exist_ok=True)
    app.run(debug=True, host="0.0.0.0", port=5000)